from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_SRC = Path(r"C:\Users\Danh\Desktop\DataMining - Copy.docx")
REPAIRED_TEMPLATE = ROOT / "outputs" / "Bao_cao_template_repaired.docx"
REPORT_PATH = ROOT / "outputs" / "Bao_cao_hoan_chinh_Bot_Detection.docx"
FIG_DIR = ROOT / "outputs" / "figures"


def repair_template() -> None:
    REPAIRED_TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    with ZipFile(TEMPLATE_SRC, "r") as zin, ZipFile(REPAIRED_TEMPLATE, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/_rels/document.xml.rels":
                text = data.decode("utf-8", errors="replace")
                text = re.sub(r'<Relationship\s+[^>]*Target="\.\./NULL"[^>]*/>', "", text)
                data = text.encode("utf-8")
            zout.writestr(item, data)


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
        caption.font.name = "Times New Roman"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        caption.font.size = Pt(12)
        caption.font.italic = True


def add_centered_paragraph(doc: Document, text: str = "", size: int = 13, bold: bool = False, space_after: int = 6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def add_para(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.0)
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.8)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run("- " + text)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(text)
    return paragraph


def add_table(doc: Document, headers: list[str], rows: list[tuple]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
        for para in table.rows[0].cells[index].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    doc.add_paragraph()
    return table


def add_picture_if_exists(doc: Document, filename: str, caption: str, width_cm: float = 15.0):
    path = FIG_DIR / filename
    if path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
        add_caption(doc, caption)
    else:
        add_para(doc, f"(Chưa tìm thấy hình {filename} trong outputs/figures.)")


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Nhấn chuột phải và chọn Update Field để cập nhật mục lục."
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def set_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def make_pipeline_figure() -> None:
    path = FIG_DIR / "system_pipeline.png"
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(12, 4.2))
    ax.axis("off")
    steps = [
        ("Cresci-2017\nusers.csv", "#d8eef0"),
        ("Tiền xử lý\nGộp dữ liệu, gán nhãn", "#f7e7b6"),
        ("Trích xuất\n20 đặc trưng hồ sơ", "#e5e1f5"),
        ("Huấn luyện\nIsolation Forest + Random Forest", "#d7ead9"),
        ("Đánh giá\nMetrics, ROC, CM", "#f3d7d7"),
        ("Triển khai\nFastAPI demo", "#d7e3f4"),
    ]
    x0, y, width, height, gap = 0.02, 0.38, 0.14, 0.28, 0.025
    for index, (label, color) in enumerate(steps):
        x = x0 + index * (width + gap)
        rect = plt.Rectangle((x, y), width, height, facecolor=color, edgecolor="#333333", linewidth=1.4)
        ax.add_patch(rect)
        ax.text(x + width / 2, y + height / 2, label, ha="center", va="center", fontsize=11, fontweight="bold")
        if index < len(steps) - 1:
            ax.annotate(
                "",
                xy=(x + width + gap * 0.75, y + height / 2),
                xytext=(x + width + gap * 0.1, y + height / 2),
                arrowprops={"arrowstyle": "->", "lw": 1.6, "color": "#333333"},
            )
    ax.text(0.5, 0.16, "Luồng xử lý từ dữ liệu thô đến mô hình dự đoán và ứng dụng demo", ha="center", fontsize=12)
    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def load_project_stats():
    features = pd.read_csv(ROOT / "data" / "processed" / "features.csv")
    train = pd.read_csv(ROOT / "data" / "processed" / "train_data.csv")
    test = pd.read_csv(ROOT / "data" / "processed" / "test_data.csv")
    metrics = json.loads((ROOT / "outputs" / "models" / "metrics.json").read_text(encoding="utf-8"))
    raw_rows = []
    for path in sorted((ROOT / "data" / "raw").glob("*/users.csv")):
        data = pd.read_csv(path, low_memory=False)
        label = "Người thật" if path.parent.name == "genuine_accounts" else "Bot"
        raw_rows.append((path.parent.name, len(data), data.shape[1], label))
    return features, train, test, metrics, raw_rows


def add_front_matter(doc: Document) -> None:
    add_centered_paragraph(doc, "TRƯỜNG ĐẠI HỌC CÔNG THƯƠNG THÀNH PHỐ HỒ CHÍ MINH", 14, True, 2)
    add_centered_paragraph(doc, "KHOA CÔNG NGHỆ THÔNG TIN", 14, True, 12)
    add_centered_paragraph(doc, "---------------------------", 13, False, 24)
    add_centered_paragraph(doc, "BÁO CÁO ĐỒ ÁN MÔN HỌC", 18, True, 8)
    add_centered_paragraph(doc, "KHAI PHÁ DỮ LIỆU", 18, True, 28)
    add_centered_paragraph(doc, "Tên đề tài", 14, True, 4)
    add_centered_paragraph(doc, "HỆ THỐNG PHÁT HIỆN TÀI KHOẢN BẤT THƯỜNG VÀ PHÂN LOẠI BOT TRÊN MẠNG XÃ HỘI", 16, True, 28)
    add_centered_paragraph(doc, "SINH VIÊN THỰC HIỆN:", 13, True, 4)
    add_centered_paragraph(doc, "1. ............................................................", 13, False, 2)
    add_centered_paragraph(doc, "2. ............................................................", 13, False, 2)
    add_centered_paragraph(doc, "3. ............................................................", 13, False, 18)
    add_centered_paragraph(doc, "GVHD: ............................................................", 13, True, 36)
    add_centered_paragraph(doc, "TP. HỒ CHÍ MINH, tháng 06 năm 2026", 13, True, 0)
    doc.add_page_break()

    doc.add_heading("Lời Cảm Ơn", level=1)
    for text in [
        "Trong quá trình thực hiện đồ án môn học Khai phá dữ liệu, nhóm chúng em đã nhận được sự định hướng và hỗ trợ từ giảng viên hướng dẫn, nhà trường và các tài liệu học thuật liên quan đến khai phá dữ liệu, học máy và phát hiện bất thường.",
        "Nhóm xin gửi lời cảm ơn chân thành đến giảng viên phụ trách học phần đã cung cấp nền tảng kiến thức về tiền xử lý dữ liệu, trích xuất đặc trưng, xây dựng mô hình, đánh giá kết quả và trình bày sản phẩm. Những kiến thức này là cơ sở để nhóm xây dựng hệ thống phát hiện tài khoản bất thường và phân loại bot trong đồ án.",
        "Nhóm cũng xin cảm ơn các thành viên đã cùng nhau thu thập dữ liệu, xây dựng notebook, huấn luyện mô hình, triển khai demo và hoàn thiện báo cáo. Do thời gian và kinh nghiệm còn hạn chế, báo cáo khó tránh khỏi thiếu sót; nhóm mong nhận được góp ý để tiếp tục hoàn thiện sản phẩm.",
    ]:
        add_para(doc, text)
    doc.add_page_break()

    doc.add_heading("Lời Cam Kết", level=1)
    for text in [
        "Nhóm chúng em cam kết báo cáo đồ án với đề tài “Hệ thống phát hiện tài khoản bất thường và phân loại bot trên mạng xã hội” là kết quả tổng hợp từ quá trình tìm hiểu tài liệu, xử lý dữ liệu, xây dựng mô hình và triển khai project trong thư mục bot-detection-project.",
        "Các số liệu thực nghiệm, biểu đồ, ma trận nhầm lẫn, chỉ số đánh giá và mô tả hệ thống trong báo cáo được lấy từ các notebook, mã nguồn và artifact đã sinh trong project. Những nội dung lý thuyết được tổng hợp từ tài liệu học thuật và tài liệu kỹ thuật có trích dẫn ở cuối báo cáo.",
        "Nhóm cam kết không sao chép nguyên văn tài liệu bên ngoài một cách thiếu trung thực. Nếu có sai sót liên quan đến nội dung học thuật hoặc kết quả thực nghiệm, nhóm xin chịu trách nhiệm và tiếp thu góp ý để chỉnh sửa.",
    ]:
        add_para(doc, text)
    doc.add_page_break()

    doc.add_heading("Mục Lục", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("Danh Mục Hình, Bảng Biểu", level=1)
    add_table(
        doc,
        ["Ký hiệu", "Tên hình/bảng"],
        [
            ("Hình 3.1", "Sơ đồ tổng quát pipeline hệ thống"),
            ("Hình 4.1", "Phân bố lớp dữ liệu"),
            ("Hình 4.2", "Quan hệ followers và friends"),
            ("Hình 4.3", "So sánh chỉ số đánh giá mô hình"),
            ("Hình 4.4", "Ma trận nhầm lẫn của hai mô hình"),
            ("Hình 4.5", "Đường cong ROC"),
            ("Hình 4.6", "Độ quan trọng đặc trưng của Random Forest"),
            ("Hình 4.7", "Phân phối điểm bất thường của Isolation Forest"),
            ("Bảng 3.1", "Các tập dữ liệu Cresci-2017 sử dụng"),
            ("Bảng 3.2", "Danh sách đặc trưng sau xử lý"),
            ("Bảng 4.1", "Môi trường và công cụ thực nghiệm"),
            ("Bảng 4.2", "Kết quả đánh giá mô hình"),
        ],
    )

    doc.add_heading("Bảng Phân Công Công Việc", level=1)
    add_table(
        doc,
        ["STT", "Họ tên", "MSSV", "Nhiệm vụ", "Mức độ hoàn thành"],
        [
            ("1", "................................", "................", "Thu thập dữ liệu, khám phá dữ liệu, viết chương tổng quan", "100%"),
            ("2", "................................", "................", "Tiền xử lý dữ liệu, trích xuất đặc trưng, xây dựng notebook", "100%"),
            ("3", "................................", "................", "Huấn luyện mô hình, đánh giá, triển khai FastAPI và hoàn thiện báo cáo", "100%"),
        ],
    )
    doc.add_page_break()


def add_opening(doc: Document) -> None:
    doc.add_heading("Mở Đầu", level=1)
    for text in [
        "Sự phát triển nhanh của mạng xã hội làm cho lượng tài khoản và nội dung do người dùng tạo ra tăng mạnh. Bên cạnh các tài khoản người thật, nhiều tài khoản tự động hoặc bán tự động được tạo ra để phát tán quảng cáo, thao túng xu hướng, lan truyền đường dẫn độc hại, tăng tương tác ảo hoặc gây nhiễu thông tin. Các tài khoản này thường được gọi chung là bot hoặc spambot.",
        "Việc nhận diện bot không chỉ dựa vào nội dung bài đăng mà còn có thể dựa trên hồ sơ tài khoản và các chỉ số hoạt động như số lượng người theo dõi, số tài khoản đang theo dõi, số bài viết, tuổi tài khoản, mức độ hoàn thiện hồ sơ và các tỷ lệ hành vi. Đây là hướng tiếp cận phù hợp với môn Khai phá dữ liệu vì yêu cầu đầy đủ các bước: thu thập dữ liệu, xử lý dữ liệu thiếu, xây dựng đặc trưng, huấn luyện mô hình, đánh giá và triển khai thử nghiệm.",
        "Đồ án này xây dựng một hệ thống phát hiện tài khoản bất thường và phân loại bot dựa trên bộ dữ liệu Cresci-2017. Hệ thống sử dụng Isolation Forest để phát hiện các hồ sơ có hành vi khác thường và Random Forest để phân loại tài khoản người thật hoặc bot. Hai kết quả được kết hợp trong giao diện demo nhằm hỗ trợ đánh giá mức rủi ro của từng tài khoản.",
    ]:
        add_para(doc, text)
    doc.add_page_break()


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("Chương 1. TỔNG QUAN", level=1)
    doc.add_heading("1.1. Đặt vấn đề", level=2)
    for text in [
        "Trên các nền tảng mạng xã hội, tài khoản bot có thể được sử dụng để tự động theo dõi nhiều người dùng, đăng bài hàng loạt, tạo tương tác giả hoặc lan truyền thông tin có chủ đích. Nếu hệ thống không có cơ chế phát hiện tự động, việc kiểm duyệt thủ công sẽ tốn nhiều thời gian, khó mở rộng và dễ bỏ sót các hành vi tinh vi.",
        "Bài toán trong đồ án được đặt ra như sau: với dữ liệu hồ sơ của một tài khoản Twitter/X, hệ thống cần trích xuất các đặc trưng số và dự đoán tài khoản đó là người thật hay bot; đồng thời đưa ra cảnh báo nếu hồ sơ có mức độ bất thường cao so với nhóm tài khoản người thật.",
    ]:
        add_para(doc, text)
    doc.add_heading("1.2. Lý do chọn đề tài", level=2)
    for text in [
        "Đề tài có ý nghĩa thực tiễn vì bot là một vấn đề phổ biến trong môi trường trực tuyến. Việc nhận diện sớm tài khoản bất thường giúp giảm nguy cơ spam, gian lận tương tác và thao túng thông tin.",
        "Đề tài cũng phù hợp với nội dung học phần Khai phá dữ liệu. Project yêu cầu xử lý dữ liệu thật, xây dựng đặc trưng, lựa chọn mô hình, đánh giá bằng nhiều độ đo và triển khai demo để kiểm thử đầu vào mới. Qua đó nhóm có thể kết nối lý thuyết học máy với một bài toán ứng dụng hoàn chỉnh.",
    ]:
        add_para(doc, text)
    doc.add_heading("1.3. Tổng quan nội dung nghiên cứu", level=2)
    add_para(doc, "Hướng nghiên cứu phát hiện bot thường chia thành ba nhóm: dựa trên nội dung bài đăng, dựa trên mạng lưới quan hệ và dựa trên thông tin hồ sơ/hành vi tổng hợp. Trong phạm vi project, nhóm tập trung vào thông tin hồ sơ người dùng từ file users.csv của Cresci-2017, vì dữ liệu này đủ để xây dựng các đặc trưng hành vi cơ bản và triển khai một demo gọn nhẹ.")
    doc.add_heading("1.3.1. Khảo sát công trình liên quan", level=3)
    for text in [
        "Bộ dữ liệu Cresci-2017 được sử dụng rộng rãi trong nghiên cứu phát hiện social spambot. Bộ dữ liệu gồm nhiều nhóm tài khoản người thật và bot, cho phép so sánh hành vi giữa các nhóm tài khoản khác nhau.",
        "Thuật toán Isolation Forest của Liu, Ting và Zhou là một phương pháp phát hiện bất thường hiệu quả vì dựa trên ý tưởng cô lập điểm dữ liệu lạ bằng cây ngẫu nhiên. Random Forest của Breiman là mô hình học có giám sát mạnh, thường được dùng cho phân loại nhị phân nhờ khả năng tổng hợp nhiều cây quyết định và giảm phương sai.",
    ]:
        add_para(doc, text)
    doc.add_heading("1.3.2. Tính mới trong phạm vi đồ án", level=3)
    for text in [
        "Project không chỉ huấn luyện một mô hình phân loại bot mà còn kết hợp thêm tầng phát hiện bất thường. Isolation Forest học chủ yếu từ tài khoản người thật để cảnh báo các hồ sơ lệch khỏi phân phối bình thường, còn Random Forest học từ dữ liệu có nhãn để đưa ra xác suất bot.",
        "Kết quả của hai mô hình được hiển thị song song trong ứng dụng FastAPI. Cách trình bày này giúp người dùng không chỉ nhận nhãn dự đoán mà còn thấy mức rủi ro tổng hợp: rủi ro thấp, cần kiểm tra hoặc rủi ro cao.",
    ]:
        add_para(doc, text)
    doc.add_heading("1.4. Mục tiêu và nhiệm vụ nghiên cứu", level=2)
    for bullet in [
        "Tải và gộp các tập users.csv của Cresci-2017, gán nhãn người thật/bot theo nguồn dữ liệu.",
        "Khám phá dữ liệu và trực quan hóa các đặc điểm nổi bật giữa tài khoản người thật và bot.",
        "Xây dựng ma trận 20 đặc trưng số từ hồ sơ tài khoản và các tỷ lệ hành vi.",
        "Huấn luyện Isolation Forest để phát hiện bất thường và Random Forest để phân loại bot.",
        "Đánh giá bằng Accuracy, Precision, Recall, F1-score, ROC-AUC và ma trận nhầm lẫn.",
        "Triển khai web demo/API cho phép nhập một tài khoản hoặc tải file CSV để dự đoán hàng loạt.",
    ]:
        add_bullet(doc, bullet)
    doc.add_heading("1.5. Bố cục đồ án", level=2)
    add_para(doc, "Báo cáo gồm năm chương. Chương 1 trình bày tổng quan và mục tiêu. Chương 2 trình bày cơ sở lý thuyết. Chương 3 mô tả phương pháp đề xuất và pipeline triển khai. Chương 4 trình bày thực nghiệm, kết quả và demo. Chương 5 tổng kết, nêu hạn chế và hướng phát triển.")


def add_chapter_2(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Chương 2. CƠ SỞ LÝ THUYẾT", level=1)
    doc.add_heading("2.1. Khai phá dữ liệu và bài toán phát hiện bot", level=2)
    for text in [
        "Khai phá dữ liệu là quá trình tìm kiếm tri thức, mẫu ẩn hoặc quy luật hữu ích từ dữ liệu lớn. Trong bài toán phát hiện bot, dữ liệu đầu vào thường chứa nhiều thuộc tính khác nhau của tài khoản như thông tin hồ sơ, chỉ số hoạt động, mối quan hệ theo dõi và hành vi đăng tải.",
        "Tài khoản bot có thể biểu hiện qua các dấu hiệu thống kê: tuổi tài khoản ngắn nhưng hoạt động dày đặc, số lượng following lớn bất thường, số followers thấp, hồ sơ thiếu mô tả, tên tài khoản chứa mẫu ký tự lặp hoặc số, hoặc tỷ lệ followers/friends khác biệt so với người thật.",
    ]:
        add_para(doc, text)
    doc.add_heading("2.2. Học có giám sát và phát hiện bất thường", level=2)
    add_para(doc, "Học có giám sát sử dụng dữ liệu đã gán nhãn để học ánh xạ từ đặc trưng đầu vào sang nhãn đầu ra. Trong project, Random Forest học từ nhãn label=0 cho người thật và label=1 cho bot. Ngược lại, phát hiện bất thường tập trung tìm các điểm dữ liệu khác biệt đáng kể so với phân phối bình thường. Isolation Forest trong project được huấn luyện trên nhóm tài khoản người thật của tập train để học vùng bình thường.")
    doc.add_heading("2.3. Bộ dữ liệu Cresci-2017", level=2)
    add_para(doc, "Cresci-2017 là bộ dữ liệu nghiên cứu về social spambots trên Twitter. Bộ dữ liệu gồm tài khoản người thật và nhiều nhóm bot khác nhau như social spambots và traditional spambots. Project sử dụng file users.csv trong từng thư mục con để xây dựng đặc trưng hồ sơ và hành vi tổng hợp, không sử dụng tweets.csv trong phạm vi triển khai hiện tại.")
    doc.add_heading("2.4. Tiền xử lý và trích xuất đặc trưng", level=2)
    for text in [
        "Tiền xử lý dữ liệu bao gồm gộp nhiều nguồn users.csv, chuẩn hóa nhãn, chuyển đổi kiểu dữ liệu, xử lý giá trị thiếu và loại bỏ các cột định danh/văn bản có số lượng giá trị quá lớn như id, name, screen_name, url, description hoặc metadata không cần thiết cho ma trận số.",
        "Feature engineering là bước biến đổi dữ liệu thô thành các đặc trưng có ý nghĩa. Project tạo thêm các đặc trưng như followers_friends_ratio, account_age_days, tweets_per_day, has_description, name_length, screen_name_digit_ratio và screen_name_has_digits. Các cột đếm bị lệch phân phối được biến đổi log1p trước khi huấn luyện để giảm ảnh hưởng của giá trị cực lớn.",
    ]:
        add_para(doc, text)
    doc.add_heading("2.5. Isolation Forest", level=2)
    for text in [
        "Isolation Forest là thuật toán phát hiện bất thường dựa trên nguyên lý: điểm dữ liệu bất thường thường dễ bị cô lập hơn điểm bình thường. Thuật toán xây dựng nhiều cây phân tách ngẫu nhiên. Một điểm càng cần ít lần phân tách để cô lập thì càng có khả năng là bất thường.",
        "Trong scikit-learn, mô hình trả về nhãn -1 cho điểm bất thường và 1 cho điểm bình thường. Project ánh xạ -1 thành lớp bot/bất thường khi đánh giá. Tham số contamination biểu diễn tỷ lệ điểm dự kiến bị xem là bất thường. Sau khi thử trên validation, project sử dụng contamination=0.25.",
        "Về trực giác, anomaly score càng cao thì tài khoản càng lệch khỏi vùng hành vi bình thường. Điểm này không thay thế xác suất bot, nhưng rất hữu ích để cảnh báo các hồ sơ cần kiểm tra thêm.",
    ]:
        add_para(doc, text)
    doc.add_heading("2.6. Random Forest", level=2)
    for text in [
        "Random Forest là mô hình tập hợp nhiều cây quyết định. Mỗi cây được huấn luyện trên một mẫu bootstrap của dữ liệu và tại mỗi nút chỉ xét một tập con đặc trưng ngẫu nhiên. Dự đoán cuối cùng được lấy bằng bỏ phiếu đa số hoặc trung bình xác suất của các cây.",
        "Ưu điểm của Random Forest là ổn định, ít nhạy với nhiễu hơn một cây đơn lẻ, xử lý tốt đặc trưng phi tuyến và cung cấp độ quan trọng đặc trưng. Trong project, RandomForestClassifier sử dụng 100 cây, random_state=42 và max_features=sqrt.",
    ]:
        add_para(doc, text)
    doc.add_heading("2.7. Các độ đo đánh giá", level=2)
    for metric in [
        "Accuracy = (TP + TN) / (TP + TN + FP + FN), thể hiện tỷ lệ dự đoán đúng trên toàn bộ tập test.",
        "Precision = TP / (TP + FP), cho biết trong các tài khoản bị dự đoán là bot thì có bao nhiêu tài khoản thật sự là bot.",
        "Recall = TP / (TP + FN), cho biết mô hình phát hiện được bao nhiêu bot trong tổng số bot thật sự.",
        "F1-score = 2 * Precision * Recall / (Precision + Recall), cân bằng giữa precision và recall.",
        "ROC-AUC đo khả năng phân tách hai lớp ở nhiều ngưỡng khác nhau. AUC càng gần 1 thì mô hình càng phân biệt tốt người thật và bot.",
    ]:
        add_bullet(doc, metric)
    doc.add_heading("2.8. FastAPI trong triển khai demo", level=2)
    add_para(doc, "FastAPI là framework Python dùng để xây dựng API nhanh và rõ ràng. Project sử dụng FastAPI để tạo giao diện web, endpoint dự đoán một tài khoản, endpoint tải CSV, endpoint tải file mẫu và endpoint tải kết quả phân tích hàng loạt. Các mô hình đã huấn luyện được lưu bằng joblib và nạp lại khi ứng dụng khởi động.")


def add_chapter_3(doc: Document, features: pd.DataFrame, raw_rows: list[tuple]) -> None:
    doc.add_page_break()
    doc.add_heading("Chương 3. PHƯƠNG PHÁP ĐỀ XUẤT", level=1)
    doc.add_heading("3.1. Mô tả bài toán", level=2)
    add_para(doc, "Đầu vào của hệ thống là hồ sơ tài khoản gồm các cột như screen_name, followers_count, friends_count, statuses_count, favourites_count, listed_count, verified, default_profile, default_profile_image, geo_enabled, protected, account_age_days hoặc created_at. Đầu ra gồm kết quả bất thường của Isolation Forest, nhãn bot/người thật của Random Forest, xác suất bot và mức rủi ro tổng hợp.")
    doc.add_heading("3.2. Sơ đồ tổng quát hệ thống", level=2)
    add_picture_if_exists(doc, "system_pipeline.png", "Hình 3.1. Sơ đồ tổng quát pipeline hệ thống", 16)
    add_para(doc, "Pipeline gồm sáu bước chính: nạp dữ liệu Cresci-2017, tiền xử lý và gán nhãn, trích xuất đặc trưng, huấn luyện hai mô hình, đánh giá kết quả và triển khai demo. Các notebook được tổ chức theo thứ tự 01_eda, 02_feature_engineering, 03_modeling và 04_evaluation để dễ tái lập.")
    doc.add_heading("3.3. Tập dữ liệu", level=2)
    add_caption(doc, "Bảng 3.1. Các tập dữ liệu Cresci-2017 sử dụng")
    add_table(doc, ["Nguồn dữ liệu", "Số dòng", "Số cột gốc", "Nhãn"], [(src, f"{rows:,}", cols, label) for src, rows, cols, label in raw_rows])
    add_para(doc, f"Sau khi gộp các tập users.csv, dữ liệu có {len(features):,} tài khoản. Trong đó có {int((features.label == 0).sum()):,} tài khoản người thật ({(features.label == 0).mean() * 100:.2f}%) và {int((features.label == 1).sum()):,} tài khoản bot ({(features.label == 1).mean() * 100:.2f}%).")
    doc.add_heading("3.4. Tiền xử lý dữ liệu", level=2)
    for text in [
        "Hàm load_cresci2017 trong src/preprocess.py duyệt qua các thư mục dữ liệu, đọc users.csv, gán nhãn theo tên nguồn và gộp thành một DataFrame chung. genuine_accounts được gán label=0, các nhóm social_spambots và traditional_spambots được gán label=1.",
        "Trong bước xây dựng đặc trưng, các cột định danh và văn bản tự do được loại bỏ khỏi ma trận huấn luyện để tránh làm mô hình học thuộc id hoặc xử lý văn bản không thuộc phạm vi project. Các giá trị vô hạn được thay bằng NaN; dữ liệu thiếu được giữ lại đến bước modeling để bộ điền giá trị trung vị chỉ fit trên tập train, tránh rò rỉ dữ liệu test.",
    ]:
        add_para(doc, text)
    doc.add_heading("3.5. Trích xuất đặc trưng", level=2)
    feature_rows = [
        ("statuses_count", "Số bài đăng/trạng thái của tài khoản"),
        ("followers_count", "Số người theo dõi tài khoản"),
        ("friends_count", "Số tài khoản mà tài khoản đang theo dõi"),
        ("favourites_count", "Số lượt thích"),
        ("listed_count", "Số danh sách mà tài khoản được đưa vào"),
        ("default_profile", "Tài khoản còn dùng hồ sơ mặc định hay không"),
        ("default_profile_image", "Tài khoản còn dùng ảnh đại diện mặc định hay không"),
        ("geo_enabled", "Có bật thông tin vị trí hay không"),
        ("profile_use_background_image", "Có dùng ảnh nền hồ sơ hay không"),
        ("utc_offset", "Thông tin múi giờ dạng số"),
        ("is_translator", "Có phải tài khoản dịch thuật hay không"),
        ("protected", "Tài khoản được bảo vệ hay không"),
        ("verified", "Tài khoản đã xác minh hay không"),
        ("followers_friends_ratio", "Tỷ lệ followers/(friends+1)"),
        ("account_age_days", "Tuổi tài khoản tính theo ngày"),
        ("tweets_per_day", "Mật độ đăng bài trung bình mỗi ngày"),
        ("has_description", "Hồ sơ có mô tả hay không"),
        ("name_length", "Độ dài tên hiển thị"),
        ("screen_name_digit_ratio", "Tỷ lệ ký tự số trong tên tài khoản"),
        ("screen_name_has_digits", "Tên tài khoản có chứa chữ số hay không"),
    ]
    add_caption(doc, "Bảng 3.2. Danh sách 20 đặc trưng sau xử lý")
    add_table(doc, ["Đặc trưng", "Ý nghĩa"], feature_rows)
    doc.add_heading("3.6. Huấn luyện mô hình", level=2)
    for text in [
        "Dữ liệu được chia train/test theo tỷ lệ 80/20 có stratify để giữ tỷ lệ nhãn ổn định. Tập train gồm 8.411 dòng và tập test gồm 2.103 dòng. Các cột đếm lệch phân phối như followers_count, friends_count, statuses_count, favourites_count và followers_friends_ratio được log-transform bằng log1p.",
        "Isolation Forest chỉ sử dụng bảy đặc trưng liên tục cốt lõi: followers_count, friends_count, followers_friends_ratio, statuses_count, favourites_count, account_age_days và tweets_per_day. Mô hình được fit trên tài khoản người thật thuộc tập train để học vùng hành vi bình thường.",
        "Random Forest sử dụng toàn bộ 20 đặc trưng sau tiền xử lý. Đây là mô hình phân loại chính vì có nhãn huấn luyện đầy đủ, trả về xác suất bot và có thể phân tích độ quan trọng đặc trưng.",
    ]:
        add_para(doc, text)
    doc.add_heading("3.7. Tổng hợp rủi ro trong demo", level=2)
    add_table(
        doc,
        ["Isolation Forest", "Random Forest", "Mức rủi ro", "Giải thích"],
        [
            ("Bất thường", "Bot", "Rủi ro cao", "Cả hai mô hình cùng cảnh báo tài khoản có dấu hiệu bot/bất thường."),
            ("Bất thường", "Người thật", "Cần kiểm tra", "Hồ sơ lệch khỏi vùng bình thường nhưng bộ phân loại chưa xác định là bot."),
            ("Bình thường", "Bot", "Cần kiểm tra", "Mô hình phân loại thấy dấu hiệu bot nhưng điểm bất thường chưa cao."),
            ("Bình thường", "Người thật", "Rủi ro thấp", "Hai mô hình đều đánh giá tài khoản có dấu hiệu bình thường."),
        ],
    )
    doc.add_heading("3.8. Cài đặt ứng dụng", level=2)
    add_para(doc, "Ứng dụng được triển khai trong src/main.py bằng FastAPI, Jinja2 template và static CSS. Khi khởi động, ứng dụng nạp imputer, scaler, Isolation Forest, Random Forest và metrics từ outputs/models. Người dùng có thể nhập nhanh một tài khoản, nhập đầy đủ thông tin hồ sơ hoặc tải file CSV để phân tích hàng loạt.")


def add_chapter_4(doc: Document, features: pd.DataFrame, train: pd.DataFrame, test: pd.DataFrame, metrics: dict) -> None:
    doc.add_page_break()
    doc.add_heading("Chương 4. THỰC NGHIỆM VÀ ĐÁNH GIÁ", level=1)
    doc.add_heading("4.1. Môi trường thực nghiệm", level=2)
    add_caption(doc, "Bảng 4.1. Môi trường và công cụ thực nghiệm")
    add_table(
        doc,
        ["Thành phần", "Công cụ/phiên bản sử dụng"],
        [
            ("Ngôn ngữ", "Python 3.x"),
            ("Xử lý dữ liệu", "pandas, numpy"),
            ("Học máy", "scikit-learn: IsolationForest, RandomForestClassifier"),
            ("Trực quan hóa", "matplotlib, seaborn"),
            ("Lưu mô hình", "joblib"),
            ("Web/API demo", "FastAPI, Uvicorn, Jinja2 Templates"),
            ("Kiểm thử", "pytest, httpx"),
        ],
    )
    doc.add_heading("4.2. Thống kê dữ liệu sau xử lý", level=2)
    add_table(
        doc,
        ["Tập dữ liệu", "Số dòng", "Người thật", "Bot", "Tỷ lệ bot"],
        [
            ("Toàn bộ features.csv", f"{len(features):,}", f"{int((features.label == 0).sum()):,}", f"{int((features.label == 1).sum()):,}", f"{(features.label == 1).mean() * 100:.2f}%"),
            ("Train", f"{len(train):,}", f"{int((train.label == 0).sum()):,}", f"{int((train.label == 1).sum()):,}", f"{(train.label == 1).mean() * 100:.2f}%"),
            ("Test", f"{len(test):,}", f"{int((test.label == 0).sum()):,}", f"{int((test.label == 1).sum()):,}", f"{(test.label == 1).mean() * 100:.2f}%"),
        ],
    )
    add_picture_if_exists(doc, "class_dist.png", "Hình 4.1. Phân bố tài khoản người thật và bot trong dữ liệu", 13.5)
    add_para(doc, "Dữ liệu có tỷ lệ bot cao hơn người thật, khoảng 66,96% bot và 33,04% người thật. Vì vậy, ngoài accuracy, báo cáo cần xem xét precision, recall, F1-score và AUC để đánh giá cân bằng hơn.")
    doc.add_heading("4.3. Khám phá quan hệ followers và friends", level=2)
    add_picture_if_exists(doc, "followers_friends_scatter.png", "Hình 4.2. Biểu đồ phân tán followers và friends theo thang log", 15)
    add_para(doc, "Biểu đồ phân tán cho thấy sự khác biệt hành vi kết nối giữa nhóm người thật và bot. Tài khoản bot thường có hồ sơ kém ổn định hơn, một số nhóm có lượng friends hoặc followers lệch mạnh so với người thật. Đây là lý do các đặc trưng liên quan đến followers_count, friends_count và followers_friends_ratio được đưa vào cả hai mô hình.")
    doc.add_heading("4.4. Tham số huấn luyện", level=2)
    add_table(
        doc,
        ["Mô hình", "Tham số chính", "Vai trò"],
        [
            ("Isolation Forest", "n_estimators=100, contamination=0.25, random_state=42", "Phát hiện tài khoản có hồ sơ bất thường"),
            ("Random Forest", "n_estimators=100, max_features=sqrt, random_state=42", "Phân loại tài khoản người thật hoặc bot"),
            ("Xử lý thiếu", "SimpleImputer/median fit trên train", "Điền giá trị thiếu, tránh leakage"),
            ("Chuẩn hóa", "StandardScaler fit trên train", "Đưa đặc trưng về cùng thang đo"),
        ],
    )
    doc.add_heading("4.5. Kết quả đánh giá mô hình", level=2)
    add_caption(doc, "Bảng 4.2. Kết quả đánh giá trên tập test")
    add_table(
        doc,
        ["Mô hình", "Accuracy", "Precision", "Recall", "F1-score", "AUC"],
        [
            ("Isolation Forest", metrics["isolation_forest"]["accuracy"], metrics["isolation_forest"]["precision"], metrics["isolation_forest"]["recall"], metrics["isolation_forest"]["f1"], metrics["isolation_forest"]["auc"]),
            ("Random Forest", metrics["random_forest"]["accuracy"], metrics["random_forest"]["precision"], metrics["random_forest"]["recall"], metrics["random_forest"]["f1"], metrics["random_forest"]["auc"]),
        ],
    )
    add_picture_if_exists(doc, "model_metrics_comparison.png", "Hình 4.3. So sánh các chỉ số đánh giá của hai mô hình", 15.5)
    add_para(doc, "Random Forest đạt Accuracy 0.9872, F1-score 0.9904 và AUC 0.9965, là mô hình phân loại chính của hệ thống. Isolation Forest đạt Recall 0.9922 đối với lớp bot/bất thường, phù hợp vai trò cảnh báo vì bỏ sót ít bot nhưng có nhiều false positive hơn so với Random Forest.")
    doc.add_heading("4.6. Ma trận nhầm lẫn", level=2)
    add_table(
        doc,
        ["Mô hình", "TN", "FP", "FN", "TP", "Nhận xét"],
        [
            ("Isolation Forest", 530, 165, 11, 1397, "Recall cao, nhưng cảnh báo nhầm một số người thật là bất thường."),
            ("Random Forest", 690, 5, 22, 1386, "Cân bằng tốt giữa phát hiện bot và giữ false positive thấp."),
        ],
    )
    add_picture_if_exists(doc, "confusion_matrices.png", "Hình 4.4. Ma trận nhầm lẫn của Isolation Forest và Random Forest", 15)
    doc.add_heading("4.7. Đường cong ROC", level=2)
    add_picture_if_exists(doc, "roc_curve.png", "Hình 4.5. Đường cong ROC của hai mô hình", 14)
    add_para(doc, "AUC của Random Forest gần 1 cho thấy khả năng phân tách hai lớp rất tốt. Isolation Forest cũng đạt AUC 0.9603, chứng tỏ điểm bất thường có tương quan mạnh với nhãn bot dù mô hình không được huấn luyện trực tiếp trên cả hai lớp như Random Forest.")
    doc.add_heading("4.8. Phân tích đặc trưng quan trọng", level=2)
    add_picture_if_exists(doc, "feature_importance.png", "Hình 4.6. Các đặc trưng quan trọng nhất của Random Forest", 14.5)
    add_table(
        doc,
        ["Xếp hạng", "Đặc trưng", "Độ quan trọng"],
        [
            ("1", "favourites_count", "0.3980"),
            ("2", "statuses_count", "0.1614"),
            ("3", "tweets_per_day", "0.1562"),
            ("4", "followers_friends_ratio", "0.0632"),
            ("5", "followers_count", "0.0608"),
            ("6", "friends_count", "0.0565"),
            ("7", "account_age_days", "0.0370"),
            ("8", "utc_offset", "0.0295"),
            ("9", "name_length", "0.0148"),
            ("10", "listed_count", "0.0096"),
        ],
    )
    add_para(doc, "Các đặc trưng liên quan đến mức độ hoạt động và tương tác như favourites_count, statuses_count, tweets_per_day và followers_friends_ratio đóng góp mạnh vào quyết định của Random Forest. Điều này phù hợp với giả thuyết rằng bot thường có mô hình hoạt động và hồ sơ khác biệt so với tài khoản người thật.")
    doc.add_heading("4.9. Phân phối điểm bất thường", level=2)
    add_picture_if_exists(doc, "anomaly_score_dist.png", "Hình 4.7. Phân phối điểm bất thường của Isolation Forest", 15)
    add_para(doc, "Phân phối anomaly score cho thấy nhóm bot có xu hướng nằm ở vùng điểm bất thường cao hơn. Tham số contamination=0.25 được chọn nhằm cân bằng giữa khả năng phát hiện và số cảnh báo nhầm, giúp mô hình phù hợp với vai trò bộ cảnh báo phụ trợ.")
    doc.add_heading("4.10. Demo ứng dụng", level=2)
    add_para(doc, "Project cung cấp web demo chạy bằng lệnh: uvicorn src.main:app --reload. Sau khi mở http://127.0.0.1:8000, người dùng có thể nhập nhanh thông tin tài khoản, nhập đầy đủ hồ sơ hoặc tải file CSV theo mẫu. File CSV đầu vào cần có các cột bắt buộc screen_name, followers_count, friends_count, statuses_count và một trong hai cột account_age_days hoặc created_at.")
    add_table(
        doc,
        ["Endpoint", "Phương thức", "Chức năng"],
        [
            ("/", "GET", "Hiển thị giao diện demo"),
            ("/predict-form", "POST", "Dự đoán một tài khoản từ form"),
            ("/upload-csv", "POST", "Dự đoán nhiều tài khoản từ file CSV"),
            ("/api/predict", "POST", "API dự đoán một tài khoản dạng JSON"),
            ("/api/upload-csv", "POST", "API phân tích CSV hàng loạt"),
            ("/download-sample-csv", "GET", "Tải file CSV mẫu"),
            ("/download-results", "GET", "Tải kết quả phân tích gần nhất"),
        ],
    )
    doc.add_heading("4.11. Nhận xét chung", level=2)
    for text in [
        "Kết quả thực nghiệm cho thấy cách kết hợp hai mô hình có ý nghĩa thực tế. Random Forest đảm nhiệm phân loại chính với độ chính xác cao, trong khi Isolation Forest cung cấp góc nhìn phát hiện bất thường độc lập. Khi cả hai mô hình cùng cảnh báo, độ tin cậy của cảnh báo rủi ro cao tăng lên.",
        "Dữ liệu hiện tại chỉ sử dụng thông tin hồ sơ tài khoản, chưa khai thác nội dung tweet, mạng lưới retweet/mention hoặc chuỗi thời gian hoạt động. Vì vậy mô hình phù hợp với demo nhận diện nhanh dựa trên metadata, nhưng chưa bao phủ toàn bộ hành vi phức tạp của bot trong môi trường thực tế.",
    ]:
        add_para(doc, text)


def add_chapter_5_and_refs(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Chương 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
    doc.add_heading("5.1. Kết luận", level=2)
    for text in [
        "Đồ án đã xây dựng được một pipeline hoàn chỉnh cho bài toán phát hiện tài khoản bất thường và phân loại bot trên mạng xã hội. Project bao gồm bước nạp dữ liệu Cresci-2017, khám phá dữ liệu, trích xuất 20 đặc trưng, huấn luyện Isolation Forest và Random Forest, đánh giá bằng nhiều độ đo và triển khai demo FastAPI.",
        "Kết quả cho thấy Random Forest đạt hiệu quả phân loại rất cao với F1-score 0.9904 và AUC 0.9965. Isolation Forest đạt recall cao, phù hợp vai trò cảnh báo bất thường. Việc kết hợp hai mô hình giúp hệ thống đưa ra đánh giá rủi ro dễ hiểu hơn so với chỉ hiển thị một nhãn phân loại.",
    ]:
        add_para(doc, text)
    doc.add_heading("5.2. Hạn chế", level=2)
    for bullet in [
        "Project mới sử dụng dữ liệu hồ sơ users.csv, chưa khai thác nội dung tweet và quan hệ mạng xã hội.",
        "Dữ liệu Cresci-2017 có thể khác với hành vi bot hiện đại, nên mô hình cần được cập nhật nếu triển khai thực tế.",
        "Một số trường dữ liệu có tỷ lệ thiếu cao như verified, is_translator hoặc default_profile_image, khiến mô hình phải phụ thuộc vào imputation.",
        "Demo hiện tập trung vào dự đoán và thống kê cơ bản, chưa có cơ chế xác thực người dùng, lưu lịch sử bằng cơ sở dữ liệu hoặc giám sát vận hành.",
    ]:
        add_bullet(doc, bullet)
    doc.add_heading("5.3. Hướng phát triển", level=2)
    for bullet in [
        "Bổ sung đặc trưng từ tweets.csv như tần suất hashtag, URL, mention, độ lặp nội dung và chuỗi thời gian đăng bài.",
        "Thử nghiệm thêm các mô hình như XGBoost, LightGBM, Logistic Regression có regularization hoặc mô hình học sâu cho dữ liệu văn bản.",
        "Tối ưu ngưỡng cảnh báo theo mục tiêu triển khai, ví dụ ưu tiên giảm false positive hoặc ưu tiên phát hiện càng nhiều bot càng tốt.",
        "Xây dựng dashboard quản trị, lưu kết quả dự đoán vào database và hỗ trợ giải thích từng dự đoán bằng SHAP hoặc permutation importance.",
        "Cập nhật dữ liệu mới hơn để kiểm tra khả năng tổng quát hóa của mô hình trên hành vi bot hiện đại.",
    ]:
        add_bullet(doc, bullet)

    doc.add_page_break()
    doc.add_heading("Tài Liệu Tham Khảo", level=1)
    refs = [
        '[1] Cresci, S., Di Pietro, R., Petrocchi, M., Spognardi, A., & Tesconi, M. "The paradigm-shift of social spambots: Evidence, theories, and tools for the arms race". WWW Companion, 2017.',
        "[2] Botometer / OSoMe. Bot Repository Datasets. https://botometer.osome.iu.edu/bot-repository/datasets.html",
        '[3] Liu, F. T., Ting, K. M., & Zhou, Z.-H. "Isolation Forest". IEEE International Conference on Data Mining, 2008.',
        '[4] Breiman, L. "Random Forests". Machine Learning, 45, 5-32, 2001.',
        "[5] scikit-learn documentation: IsolationForest and RandomForestClassifier. https://scikit-learn.org/stable/",
        "[6] FastAPI documentation. https://fastapi.tiangolo.com/",
        "[7] Project source code: bot-detection-project, các file src/preprocess.py, src/features.py, src/models.py, src/main.py và notebooks 01-04.",
    ]
    for ref in refs:
        paragraph = doc.add_paragraph(ref)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(-0.6)
        paragraph.paragraph_format.left_indent = Cm(0.6)

    doc.add_page_break()
    doc.add_heading("Phụ Lục A. Hướng Dẫn Chạy Project", level=1)
    add_para(doc, "Cài đặt thư viện: pip install -r requirements.txt")
    add_para(doc, "Chạy notebook theo thứ tự: notebooks/01_eda.ipynb, notebooks/02_feature_engineering.ipynb, notebooks/03_modeling.ipynb và notebooks/04_evaluation.ipynb.")
    add_para(doc, "Chạy web demo: uvicorn src.main:app --reload, sau đó mở http://127.0.0.1:8000 trên trình duyệt.")
    add_para(doc, "Chạy kiểm thử API: pytest -q.")


def build_report() -> None:
    repair_template()
    make_pipeline_figure()
    features, train, test, metrics, raw_rows = load_project_stats()

    doc = Document(REPAIRED_TEMPLATE)
    clear_document_body(doc)
    set_document_style(doc)
    set_update_fields(doc)

    add_front_matter(doc)
    add_opening(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc, features, raw_rows)
    add_chapter_4(doc, features, train, test, metrics)
    add_chapter_5_and_refs(doc)

    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    build_report()
